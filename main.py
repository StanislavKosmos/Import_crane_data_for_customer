from docx import Document
import openpyxl as xl
from docx.shared import Cm, Mm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import add_pict_docx as add_pict
import os
import sys 


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def set_col_widths(table, widths):
    """
    Set a column width for table
    """
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_tabl_txt(table, text, align, row, col, txt_height):
    """
    Запись в ячейку таблицы текста с заданными параметрами
    :param table: таблица
    :param text: текст, строковая или другая переменная
    :param align_txt: выравнивание текста (CENTER, LEFT, RIGHT)
    :param row: номер строки
    :param col: номер столбца
    :param txt_height: высота текста в Pt
    """

    p = table.cell(row, col).paragraphs[0]
    if align == 'C':
        aligh_t = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'L':
        aligh_t = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'R':
        aligh_t = WD_ALIGN_PARAGRAPH.RIGHT
    p.alignment = aligh_t
    run_el = p.add_run(text)
    run_el.font.size = Pt(txt_height)


def write_crane_inf(self, sheet, txtHeight):
    """
    create table with crane info from xlsx sheet
    """
    table = self.add_table(rows=1, cols=2, style='Style_without_border')
    # зададим нужную ширину столбцов таблицы
    widths_t = [Mm(60), Mm(80)]
    i = 0
    row_num = 3
    item = 'Starting'
    while item is not None:
        if i != 0:
            table.add_row().cells
        j = 0
        for col_num in range(11, 13):
            item = sheet.cell(row=row_num, column=col_num).value
            if type(item) == float or type(item) == int:
                item = str(round(item, 3))
            cell = table.rows[i].cells[j]
            run_elem = cell.paragraphs[0].add_run(item)
            run_elem.font.size = Pt(txtHeight)
            j += 1
        row_num += 1
        i += 1
        item = sheet.cell(row=row_num, column=col_num).value
    set_col_widths(table, widths_t)
    return table


def create_table_factors(self, sheet, txtHeight):
    """
    create table with factors (dynamic etc.) from xlsx sheet
    """
    table = self.add_table(rows=10, cols=3, style='MyStyle')
    widths = [Mm(15), Mm(15), Mm(100)]
    set_col_widths(table, widths)
    t_alignment = ['C', 'C', 'L']
    txt = ['Коэффициент', 'φ', 'φ', 'φ', 'φ', 'φ', 'φ', 'γ', 'γ', 'γ']
    sub_txt = ['', '1', '2', '5 тел', '5 кран', '6', '7', 'гр', 'кран', 'n']
    i = 0
    for row_num in range(1, 11):
        j = 0
        # пишем первый столбец отдельно
        p = table.cell(i, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_el = p.add_run(txt[i])
        run_el.font.size = Pt(txtHeight)
        subtxt = p.add_run(sub_txt[i])
        subtxt.font.size = Pt(txtHeight)
        subtxt.font.subscript = True
        # пишем другие столбцы
        for col_num in range(2, 4):
            item = sheet.cell(row=row_num, column=col_num).value
            if type(item) == float:
                item = str(round(item, 3))
            add_tabl_txt(table, text=item, align=t_alignment[j+1], row=i, col=j+1, txt_height=txtHeight)
            j += 1
        i += 1
    return table


def create_table_vert_force(self, sheet, numWheels, txtHeight):
    """
    create vertical forces of crane wheels from xlsx sheet
    """
    table = self.add_table(rows=12, cols=numWheels+1, style='MyStyle')
    t_widths = [Mm(75)]
    for _ in range(numWheels):
        t_widths.append(Mm(12.5))
    set_col_widths(table, t_widths)
    forces = []
    for i in range(20, 31):
        table1 = []
        for j in range(1, numWheels+2):
            item = sheet.cell(row=i, column=j).value
            if type(item) == float:
                item = round(item, 1)
            table1.append(item)
        forces.append(table1)
    # объединим нужные нам ячейки для необходимого вида таблицы
    table.cell(0, 0).merge(table.cell(1, 0))
    cell = table.cell(0, 1)
    for i in range(2, len(forces[0])):
        cell.merge(table.cell(0, i))
    # запишем в ячейки сначала все силы
    for row_num in range(1, len(forces)):
        for col_num in range(1, len(forces[row_num])):
            add_tabl_txt(table, text=str(forces[row_num][col_num]), align='C',
                         row=row_num+1, col=col_num, txt_height=txtHeight)
    # запись текста в левую верхнюю ячейку таблицы
    add_tabl_txt(table, text='Нагрузки', align='C', row=0, col=0, txt_height=txtHeight)
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_tabl_txt(table, text='Колёса', align='C', row=0, col=1, txt_height=txtHeight)
    for col_num in range(1, len(forces[0])):
        add_tabl_txt(table, text=str(forces[0][col_num]), align='C', row=1, col=col_num, txt_height=txtHeight)
    for row_num in range(1, len(forces)):
        p = table.cell(row_num+1, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_el = p.add_run(str(forces[row_num][0])+' F')
        run_el.font.size = Pt(txtHeight)
        subtxt = p.add_run('z i-j')
        subtxt.font.size = Pt(txtHeight)
        subtxt.font.subscript = True
    return table


def create_table_horiz_force(self, sheet, numWheels, txtHeight):
    """
    create vertical forces of crane wheels from xlsx sheet
    """
    table = self.add_table(rows=12, cols=numWheels+1, style='MyStyle')
    t_widths = [Mm(75)]
    for _ in range(numWheels):
        t_widths.append(Mm(12.5))
    set_col_widths(table, t_widths)
    # считаем с Excel таблицу горизонтальных усилий
    forces = []
    for i in range(50, 61):
        table1 = []
        for j in range(1, numWheels+2):
            item = sheet.cell(row=i, column=j).value
            if type(item) == float:
                item = round(item, 1)
            table1.append(item)
        forces.append(table1)
    # считаем с Excel боковые контактные усилия, возникающие при перекосе крана
    SideForces = []
    for i in range(61, 63):
        item = sheet.cell(row=i, column=2).value
        if type(item) == float:
            item = str(round(item, 1))
        SideForces.append(item)
    # объединим нужные нам ячейки для необходимого вида таблицы
    table.cell(0, 0).merge(table.cell(1, 0))
    cell = table.cell(0, 1)
    for i in range(2, len(forces[0])):
        cell.merge(table.cell(0, i))
    # запись текста в верхние ячейки таблицы
    add_tabl_txt(table, text='Нагрузки', align='C', row=0, col=0, txt_height=txtHeight)
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_tabl_txt(table, text='Колёса', align='C', row=0, col=1, txt_height=txtHeight)
    # запишем номера колёс
    for col_num in range(1, len(forces[0])):
        txtForCell = str(forces[0][col_num])
        add_tabl_txt(table, text=txtForCell, align='C', row=1, col=col_num, txt_height=txtHeight)
    sub_txt_hor_f = ['y i-j', 'x i-j', 'x i-j', 'x i-j', 'y i-j', 'y i-j',
                     'x i-j', 'x i-j', 'x i-j', 'x i-j']
    for row_num in range(1, len(forces)):
        p = table.cell(row_num+1, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        txtForCell = str(forces[row_num][0])
        run_el = p.add_run(txtForCell + ' F')
        run_el.font.size = Pt(txtHeight)
        subtxt = p.add_run(sub_txt_hor_f[row_num-1])
        subtxt.font.size = Pt(txtHeight)
        subtxt.font.subscript = True
        if txtForCell == 'Перекос крана (ПТ1), кН':
            txt = p.add_run(', S')
            txt.font.size = Pt(txtHeight)
            subtxt = p.add_run('1,2')
            subtxt.font.size = Pt(txtHeight)
            subtxt.font.subscript = True
            txt = p.add_run('=')
            txt.font.size = Pt(txtHeight)
            SideForcetxt = p.add_run(str(SideForces[0]))
            SideForcetxt.font.size = Pt(txtHeight)
            txt = p.add_run(' кН')
            txt.font.size = Pt(txtHeight)
        elif txtForCell == 'Перекос крана (ПТ2), кН':
            txt = p.add_run(', S')
            txt.font.size = Pt(txtHeight)
            subtxt = p.add_run('2,2')
            subtxt.font.size = Pt(txtHeight)
            subtxt.font.subscript = True
            txt = p.add_run('=')
            txt.font.size = Pt(txtHeight)
            SideForcetxt = p.add_run(str(SideForces[1]))
            SideForcetxt.font.size = Pt(txtHeight)
            txt = p.add_run(' кН')
            txt.font.size = Pt(txtHeight)
    # запишем в ячейки все силы с необходимым объединением
    for row_num in range(1, len(forces)):
        row_tabl = row_num + 1
        a = forces[row_num]
        try:
            digit = a.count(None)
            if digit == numWheels - 2:
                cell = table.cell(row_tabl, 1)
                for j in range(2, int(numWheels/2)+1):
                    cell.merge(table.cell(row_tabl, j))
                cell = table.cell(row_tabl, int(numWheels/2)+1)
                for j in range(int(numWheels/2)+2, numWheels+1):
                    cell.merge(table.cell(row_tabl, j))
            elif digit > numWheels - 2:
                info_txt = ''
                cell = table.cell(row_tabl, 1)
                for i in range(2, len(forces[0])):
                    cell.merge(table.cell(row_tabl, i))
            for col_num in range(1, len(forces[row_num])):
                item = forces[row_num][col_num]
                if item is None:
                    continue
                else:
                    if digit == numWheels - 2:
                        num_rail = str(forces[0][col_num][0])
                        info_txt = f'R{num_rail} = '
                    else:
                        info_txt = ''
                    add_tabl_txt(table, text=info_txt+str(item), align='C',
                                 row=row_tabl, col=col_num, txt_height=txtHeight)
        except ValueError:
            for col_num in range(1, len(forces[row_num])):
                item = str(forces[row_num][col_num])
                add_tabl_txt(table, text=item, align='C', row=row_tabl, col=col_num, txt_height=txtHeight)
    return table


def import_data_for_customer(self, sheetxl, TxtHeight):
    """
    import data about crane for the customer
    """
    table = self.add_table(rows=1, cols=2, style='Normal Table Style')
    i = 1
    for colum in range(1, 11, 3):
        row = 1
        celltxt = 'Start'
        while celltxt is not None:
            for j in range(1, 3):
                celltxt = sheetxl.cell(row=row, column=colum+j-1).value
                if j == 2 and celltxt is None:
                    table.cell(i-1, 0).merge(table.cell(i-1, 1))
                    continue
                cell = table.rows[i-1].cells[j-1]
                table.cell(i-1, j-1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run_elem = cell.paragraphs[0].add_run(str(celltxt))
                if j == 1 and sheetxl.cell(row=row, column=colum+1).value is None:
                    run_elem.bold = True
                run_elem.font.size = Pt(TxtHeight)
                run_elem.font.name = 'Times New Roman'
            row += 1
            i += 1
            table.add_row().cells
            celltxt = sheetxl.cell(row=row, column=colum).value
    row_last = table.rows[len(table.rows)-1]
    row_last._element.getparent().remove(row_last._element)
    return table


def importCraneInfo2Cust():
    path = os.getcwd()
    # download XLSM document with data from MathCad
    file_xls = 'Выгрузка для менеджеров.xlsm'
    wb = xl.load_workbook(file_xls)
    sheetxl = wb['Лист3']
    name_doc = wb['Лист5']['L1'].value
    os.chdir(path+'\Other files')
    doc = Document('template_A4_port_2b.docx')
    # start creating the table for customer
    section = doc.sections[-1]
    # field sizes
    section.footer_distance = Mm(5)
    delete_paragraph(doc.paragraphs[0])
    head1 = doc.add_paragraph()
    head1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head1.add_run('Приложение А. Выгрузка данных по крану')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    head1.paragraph_format.space_after = Pt(10)
    # пробел между таблицей и дальнейшим параграфом
    doc.add_paragraph()
    # создадим таблицу с данными крана для заказчика
    DataForCustomer = import_data_for_customer(doc, sheetxl, TxtHeight=10)
    DataForCustomer.alignment = WD_TABLE_ALIGNMENT.LEFT
    # заполним в рамке через нижний колонтитул наименование документа
    table_footer = section.footer.add_table(rows=1, cols=1, width=Cm(10))
    table_footer.alignment = WD_TABLE_ALIGNMENT.RIGHT
    p = table_footer.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_el = p.add_run(name_doc)
    run_el.font.name = 'Times New Roman'
    run_el.font.size = Pt(20)
    table_footer.cell(0, 0).width = Cm(8.8)
    os.chdir(os.pardir)
    doc.save('Приложение А. Выгрузка данных по крану.docx')


def importWheelForces():
    name_file = 'Приложение Б. Нагрузки на крановые колёса.docx'
    path = os.getcwd()
    # подгрузим XLSX документ с выгрузкой из MathCad
    file_xls = 'Выгрузка для менеджеров.xlsm'
    wb = xl.load_workbook(file_xls)
    sheet = wb['Лист5']
    name_doc = sheet['L1'].value
    numWheels = sheet['L2'].value
    # откроем исходный документ и начнем заполнять шапку документа
    os.chdir(path + '/Other files')
    doc = Document('template_A4_land_2b.docx')
    section = doc.sections[-1]
    # размеры полей
    section.footer_distance = Mm(9)
    delete_paragraph(doc.paragraphs[0])
    head1 = doc.add_paragraph()
    head1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head1.add_run('Приложение Б. Нагрузки на крановые колёса')
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.bold = True
    head1.paragraph_format.space_after = Pt(10)
    # вставим картинку в нужное нам место
    add_pict.add_float_picture(head1, str(numWheels)+' wheel forces.png', width=Mm(108), pos_x=Mm(175), pos_y=Mm(17))
    # вставим таблицы информации о кране
    table_cr_info = write_crane_inf(doc, sheet, txtHeight=10)
    table_cr_info.alignment = WD_TABLE_ALIGNMENT.LEFT
    # пробел между таблицей и дальнейшим параграфом
    doc.add_paragraph()
    # наименование таблицы 1
    para = doc.add_paragraph()
    para_fmt = para.paragraph_format
    para_fmt.left_indent = Mm(10)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_elem = para.add_run('Таблица 1. Коэффициенты в соответствии с ГОСТ 32579.1-2013, используемые в расчёте крана')
    run_elem.font.size = Pt(8)
    # создадим таблицу динамических коэффициентов
    table_factors = create_table_factors(doc, sheet, txtHeight=8)
    table_factors.alignment = WD_TABLE_ALIGNMENT.LEFT
    # пробел между таблицей и дальнейшим параграфом
    doc.add_paragraph()
    # пишем заголовок таблицы 2
    para = doc.add_paragraph()
    para_fmt = para.paragraph_format
    para_fmt.left_indent = Mm(10)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_elem = para.add_run('Таблица 2. Вертикальные нагрузки на колёса (даны без учёта коэффициентов надежности, '
                            'ответственности)')
    run_elem.font.size = Pt(8)
    create_table_vert_force(doc, sheet, numWheels, txtHeight=8)
    # пробел между таблицей и дальнейшим параграфом
    doc.add_paragraph()
    # пишем заголовок таблицы 3
    para = doc.add_paragraph()
    para_fmt = para.paragraph_format
    para_fmt.left_indent = Mm(10)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_elem = para.add_run('Таблица 3. Горизонтальные нагрузки на колёса (даны без учёта коэффициентов надежности, '
                            'ответственности)')
    run_elem.font.size = Pt(8)
    create_table_horiz_force(doc, sheet, numWheels, txtHeight=8)
    # запишем пояснение под таблицей насчёт указанных обозначений
    para = doc.add_paragraph()
    para_fmt = para.paragraph_format
    para_fmt.left_indent = Mm(10)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_elem = para.add_run('ПТ1 и ПТ2 - положение тельфера у рельса 1 и 2 соответственно, '
                            'R1 и R2 - суммарная продольная нагрузка на рельсе 1 и 2 соответственно, S')
    run_elem.font.size = Pt(8)
    subtxt = para.add_run('1,2')
    subtxt.font.size = Pt(8)
    subtxt.font.subscript = True
    txt = para.add_run(' и S')
    txt.font.size = Pt(8)
    subtxt = para.add_run('2,2')
    subtxt.font.size = Pt(8)
    subtxt.font.subscript = True
    txt = para.add_run(' - боковое контактное усилие, когда положение тельфера у рельса 1 и 2 соответственно')
    txt.font.size = Pt(8)
    # запись текста в основную надпись документа
    table_footer = section.footer.add_table(rows=1, cols=1, width=Cm(10))
    table_footer.alignment = WD_TABLE_ALIGNMENT.RIGHT
    add_tabl_txt(table_footer, text=name_doc, align='L', row=0, col=0, txt_height=20)
    table_footer.cell(0, 0).width = Cm(8.5)
    os.chdir(os.pardir)
    doc.save(name_file)
    sys.stderr = open(path+"/Other files/consoleoutput.log", "w")
    convert(name_file)
    os.remove(name_file)


if __name__ == '__main__':
    importCraneInfo2Cust()
    importWheelForces()
